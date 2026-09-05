"""Integration tests for Phase 2 office routes (Word template + PDF ops).

Covers the 6 new endpoints added in Task 8:
- POST /api/v1/office/word/analyze-template
- POST /api/v1/office/word/fill-template
- POST /api/v1/office/pdf/read
- POST /api/v1/office/pdf/generate
- POST /api/v1/office/pdf/read-form
- POST /api/v1/office/pdf/fill-form

Tests call endpoint functions directly (matching the existing pattern in
test_office_routes.py). Each test verifies response shape and error handling.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from fastapi.testclient import TestClient
from pydantic import ValidationError

from backend.main import app
from backend.office.models import (
    PdfFormFillRequest,
    PdfFormReadRequest,
    PdfGenerateRequest,
    PdfPageSpec,
    PdfReadRequest,
    WordTemplateAnalyzeRequest,
    WordTemplateFillRequest,
)


@pytest.fixture()
def workspace(tmp_path: Path) -> Path:
    """Per-test scratch workspace."""
    ws = tmp_path / "workspace"
    ws.mkdir()
    return ws


@pytest.fixture()
def http_client() -> TestClient:
    """Authenticated client exercising the real FastAPI route stack."""
    return TestClient(
        app,
        headers={"X-Sage-Local-Authorization": "Bearer test-local-auth-token"},
    )


# ──────────────────────────────────────────────────────────────────────
# Helpers to create real test files
# ──────────────────────────────────────────────────────────────────────


def _make_word_template(workspace: Path) -> Path:
    """Create a minimal Word template with {{name}} and {{title}} placeholders."""
    from docx import Document

    path = workspace / "template.docx"
    doc = Document()
    doc.add_paragraph("Hello {{ name }}!")
    doc.add_paragraph("Your title is {{ title }}.")
    doc.save(str(path))
    return path


def _make_pdf(workspace: Path) -> Path:
    """Create a minimal PDF with one page of text."""
    from reportlab.pdfgen import canvas as rl_canvas

    path = workspace / "input.pdf"
    c = rl_canvas.Canvas(str(path))
    c.drawString(72, 700, "Hello PDF world")
    c.save()
    return path


def _make_pdf_form(workspace: Path) -> Path:
    """Create a PDF with two AcroForm text fields."""
    import pymupdf

    path = workspace / "form.pdf"
    doc = pymupdf.open()
    page = doc.new_page()

    for i, field_name in enumerate(["name", "email"]):
        widget = pymupdf.Widget()
        widget.field_name = field_name
        widget.field_type = pymupdf.PDF_WIDGET_TYPE_TEXT
        widget.rect = pymupdf.Rect(72, 700 - i * 50, 300, 720 - i * 50)
        widget.field_value = ""
        page.add_widget(widget)

    doc.save(str(path))
    doc.close()
    return path


# ──────────────────────────────────────────────────────────────────────
# Word template endpoints
# ──────────────────────────────────────────────────────────────────────


def test_analyze_word_template_endpoint(workspace: Path) -> None:
    """POST /office/word/analyze-template returns placeholders."""
    from backend.api.office_routes import analyze_word_template_endpoint

    template_path = _make_word_template(workspace)

    result = analyze_word_template_endpoint(
        WordTemplateAnalyzeRequest(
            workspace_path=str(workspace),
            template_path=str(template_path),
        )
    )

    assert result.file_path == str(template_path)
    assert isinstance(result.placeholders, list)
    placeholder_names = {p.name.strip() for p in result.placeholders}
    assert "name" in placeholder_names
    assert "title" in placeholder_names
    assert result.summary is not None
    assert result.summary.doc_type == "word"


def test_analyze_word_template_missing_file(workspace: Path) -> None:
    """analyze_word_template raises OfficePathError for missing file (route layer)."""
    from backend.api.office_routes import analyze_word_template_endpoint
    from backend.office.errors import OfficePathError

    with pytest.raises(OfficePathError):
        analyze_word_template_endpoint(
            WordTemplateAnalyzeRequest(
                workspace_path=str(workspace),
                template_path=str(workspace / "nonexistent.docx"),
            )
        )


def test_fill_word_template_endpoint(workspace: Path) -> None:
    """POST /office/word/fill-template fills placeholders and writes output."""
    from backend.api.office_routes import fill_word_template_endpoint

    template_path = _make_word_template(workspace)

    result = fill_word_template_endpoint(
        WordTemplateFillRequest(
            workspace_path=str(workspace),
            template_path=str(template_path),
            output_filename="filled.docx",
            data={"name": "Alice", "title": "Engineer"},
        )
    )

    assert result.output_path.endswith("filled.docx")
    assert result.filename == "filled.docx"
    assert result.file_size_bytes > 0
    assert result.filled_count == 2
    assert Path(result.output_path).is_file()


def test_fill_word_template_missing_template(workspace: Path) -> None:
    """fill_word_template raises OfficeFileNotFoundError for missing template."""
    from backend.api.office_routes import fill_word_template_endpoint
    from backend.office.errors import OfficeFileNotFoundError

    with pytest.raises(OfficeFileNotFoundError):
        fill_word_template_endpoint(
            WordTemplateFillRequest(
                workspace_path=str(workspace),
                template_path=str(workspace / "nonexistent.docx"),
                output_filename="out.docx",
                data={"name": "Alice"},
            )
        )


# ──────────────────────────────────────────────────────────────────────
# PDF endpoints
# ──────────────────────────────────────────────────────────────────────


def test_read_pdf_endpoint(workspace: Path) -> None:
    """POST /office/pdf/read returns pages with text."""
    from backend.api.office_routes import read_pdf_endpoint

    pdf_path = _make_pdf(workspace)

    result = read_pdf_endpoint(
        PdfReadRequest(
            workspace_path=str(workspace),
            file_path=str(pdf_path),
        )
    )

    assert result.summary is not None
    assert result.summary.doc_type == "pdf"
    assert len(result.pages) == 1
    assert "Hello PDF world" in result.pages[0].text


def test_read_pdf_missing_file(workspace: Path) -> None:
    """read_pdf raises OfficePathError for missing file (route layer)."""
    from backend.api.office_routes import read_pdf_endpoint
    from backend.office.errors import OfficePathError

    with pytest.raises(OfficePathError):
        read_pdf_endpoint(
            PdfReadRequest(
                workspace_path=str(workspace),
                file_path=str(workspace / "nonexistent.pdf"),
            )
        )


def test_generate_pdf_endpoint(workspace: Path) -> None:
    """POST /office/pdf/generate creates a PDF file."""
    from backend.api.office_routes import generate_pdf_endpoint

    result = generate_pdf_endpoint(
        PdfGenerateRequest(
            workspace_path=str(workspace),
            filename="output.pdf",
            pages=[
                PdfPageSpec(title="Test Page", paragraphs=["Hello", "World"]),
            ],
        )
    )

    assert result.filename == "output.pdf"
    assert result.file_size_bytes > 0
    assert result.page_count >= 1
    assert Path(result.output_path).is_file()


def test_generate_pdf_invalid_filename(workspace: Path) -> None:
    """generate_pdf rejects absolute path filenames."""
    from backend.api.office_routes import generate_pdf_endpoint
    from backend.office.errors import OfficePdfGenerateError

    with pytest.raises(OfficePdfGenerateError):
        generate_pdf_endpoint(
            PdfGenerateRequest(
                workspace_path=str(workspace),
                filename="/etc/passwd",
                pages=[PdfPageSpec(paragraphs=["x"])],
            )
        )


def test_read_pdf_form_endpoint(workspace: Path) -> None:
    """POST /office/pdf/read-form returns form fields."""
    from backend.api.office_routes import read_pdf_form_endpoint

    form_path = _make_pdf_form(workspace)

    result = read_pdf_form_endpoint(
        PdfFormReadRequest(
            workspace_path=str(workspace),
            file_path=str(form_path),
        )
    )

    assert result.file_path == str(form_path)
    assert isinstance(result.fields, list)
    field_names = {f.name for f in result.fields}
    assert "name" in field_names
    assert "email" in field_names


def test_read_pdf_form_missing_file(workspace: Path) -> None:
    """read_pdf_form raises OfficePathError for missing file (route layer)."""
    from backend.api.office_routes import read_pdf_form_endpoint
    from backend.office.errors import OfficePathError

    with pytest.raises(OfficePathError):
        read_pdf_form_endpoint(
            PdfFormReadRequest(
                workspace_path=str(workspace),
                file_path=str(workspace / "nonexistent.pdf"),
            )
        )


def test_fill_pdf_form_endpoint(workspace: Path) -> None:
    """POST /office/pdf/fill-form fills fields and writes output."""
    from backend.api.office_routes import fill_pdf_form_endpoint

    form_path = _make_pdf_form(workspace)

    result = fill_pdf_form_endpoint(
        PdfFormFillRequest(
            workspace_path=str(workspace),
            template_path=str(form_path),
            output_filename="filled_form.pdf",
            data={"name": "Alice", "email": "alice@example.com"},
        )
    )

    assert result.filename == "filled_form.pdf"
    assert result.file_size_bytes > 0
    assert result.filled_count >= 1
    assert Path(result.output_path).is_file()


def test_fill_pdf_form_missing_template(workspace: Path) -> None:
    """fill_pdf_form raises OfficeFileNotFoundError for missing template."""
    from backend.api.office_routes import fill_pdf_form_endpoint
    from backend.office.errors import OfficeFileNotFoundError

    with pytest.raises(OfficeFileNotFoundError):
        fill_pdf_form_endpoint(
            PdfFormFillRequest(
                workspace_path=str(workspace),
                template_path=str(workspace / "nonexistent.pdf"),
                output_filename="out.pdf",
                data={"name": "Alice"},
            )
        )


# ──────────────────────────────────────────────────────────────────────
# Request model validation (extra fields forbidden)
# ──────────────────────────────────────────────────────────────────────


def test_analyze_request_rejects_extra_fields() -> None:
    """WordTemplateAnalyzeRequest rejects unknown fields (extra='forbid')."""
    with pytest.raises(ValidationError):
        WordTemplateAnalyzeRequest(
            workspace_path="/tmp/ws",
            template_path="/tmp/ws/t.docx",
            bogus="not allowed",
        )


def test_pdf_read_request_rejects_extra_fields() -> None:
    """PdfReadRequest rejects unknown fields."""
    with pytest.raises(ValidationError):
        PdfReadRequest(
            workspace_path="/tmp/ws",
            file_path="/tmp/ws/f.pdf",
            bogus="not allowed",
        )


def test_pdf_form_read_request_rejects_extra_fields() -> None:
    """PdfFormReadRequest rejects unknown fields."""
    with pytest.raises(ValidationError):
        PdfFormReadRequest(
            workspace_path="/tmp/ws",
            file_path="/tmp/ws/f.pdf",
            bogus="not allowed",
        )


# ──────────────────────────────────────────────────────────────────────
# Real FastAPI HTTP route coverage
# ──────────────────────────────────────────────────────────────────────


def test_http_analyze_word_template_route(
    workspace: Path, http_client: TestClient
) -> None:
    """The mounted analyze route returns its concrete response schema."""
    template_path = _make_word_template(workspace)
    response = http_client.post(
        "/api/v1/office/word/analyze-template",
        json={
            "workspace_path": str(workspace),
            "template_path": str(template_path),
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert {item["name"].strip() for item in body["placeholders"]} == {
        "name",
        "title",
    }
    assert body["summary"]["doc_type"] == "word"


def test_http_fill_word_template_route_checks_output_content(
    workspace: Path, http_client: TestClient
) -> None:
    """The mounted fill route writes a DOCX containing substituted values."""
    from docx import Document

    template_path = _make_word_template(workspace)
    response = http_client.post(
        "/api/v1/office/word/fill-template",
        json={
            "workspace_path": str(workspace),
            "template_path": str(template_path),
            "output_filename": "http-filled.docx",
            "data": {"name": "Bob", "title": "Reviewer"},
        },
    )

    assert response.status_code == 200
    body = response.json()
    output_path = Path(body["output_path"])
    assert output_path.is_file()
    output_text = "\n".join(
        paragraph.text for paragraph in Document(str(output_path)).paragraphs
    )
    assert "Hello Bob!" in output_text
    assert "Your title is Reviewer." in output_text


def test_http_read_pdf_route(http_client: TestClient, workspace: Path) -> None:
    """The mounted PDF read route returns page text."""
    pdf_path = _make_pdf(workspace)
    response = http_client.post(
        "/api/v1/office/pdf/read",
        json={"workspace_path": str(workspace), "file_path": str(pdf_path)},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["summary"]["doc_type"] == "pdf"
    assert body["pages"][0]["text"].strip() == "Hello PDF world"


def test_http_generate_pdf_route_checks_output_content(
    http_client: TestClient, workspace: Path
) -> None:
    """The mounted PDF generate route writes a readable PDF."""
    import pymupdf

    response = http_client.post(
        "/api/v1/office/pdf/generate",
        json={
            "workspace_path": str(workspace),
            "filename": "http-output.pdf",
            "pages": [{"title": "HTTP", "paragraphs": ["Generated content"]}],
        },
    )

    assert response.status_code == 200
    output_path = Path(response.json()["output_path"])
    assert output_path.is_file()
    with pymupdf.open(str(output_path)) as document:
        assert "Generated content" in "\n".join(page.get_text() for page in document)


def test_http_read_pdf_form_route(
    http_client: TestClient, workspace: Path
) -> None:
    """The mounted PDF form read route returns concrete field metadata."""
    form_path = _make_pdf_form(workspace)
    response = http_client.post(
        "/api/v1/office/pdf/read-form",
        json={"workspace_path": str(workspace), "file_path": str(form_path)},
    )

    assert response.status_code == 200
    assert {field["name"] for field in response.json()["fields"]} == {
        "name",
        "email",
    }


def test_http_fill_pdf_form_route_checks_filled_values(
    http_client: TestClient, workspace: Path
) -> None:
    """The mounted PDF form fill route persists submitted widget values."""
    import pymupdf

    form_path = _make_pdf_form(workspace)
    response = http_client.post(
        "/api/v1/office/pdf/fill-form",
        json={
            "workspace_path": str(workspace),
            "template_path": str(form_path),
            "output_filename": "http-filled-form.pdf",
            "data": {"name": "Carol", "email": "carol@example.com"},
        },
    )

    assert response.status_code == 200
    output_path = Path(response.json()["output_path"])
    assert output_path.is_file()
    with pymupdf.open(str(output_path)) as document:
        values = {
            widget.field_name: widget.field_value
            for page in document
            if page.widgets()
            for widget in page.widgets()
        }
    assert values["name"] == "Carol"
    assert values["email"] == "carol@example.com"


def test_http_office_error_is_structured(
    http_client: TestClient, workspace: Path
) -> None:
    """OfficeError handlers expose a 400 structured path error response."""
    response = http_client.post(
        "/api/v1/office/pdf/read",
        json={
            "workspace_path": str(workspace),
            "file_path": str(workspace.parent / "outside.pdf"),
        },
    )

    assert response.status_code == 400
    assert response.json()["error_type"] == "OfficePathError"
    assert "message" in response.json()


def test_http_write_route_rejects_traversal_name(
    http_client: TestClient, workspace: Path
) -> None:
    """The mounted PDF generate route rejects traversal output names."""
    response = http_client.post(
        "/api/v1/office/pdf/generate",
        json={
            "workspace_path": str(workspace),
            "filename": "../escape.pdf",
            "pages": [{"paragraphs": ["blocked"]}],
        },
    )

    assert response.status_code == 500
    assert response.json()["error_type"] == "OfficePdfGenerateError"


def test_http_routes_require_local_auth(workspace: Path) -> None:
    """The real mounted route is protected by LocalAuthMiddleware."""
    template_path = _make_word_template(workspace)
    unauthenticated = TestClient(
        app,
        headers={"X-Sage-Local-Authorization": "Bearer wrong-token"},
    )
    response = unauthenticated.post(
        "/api/v1/office/word/analyze-template",
        json={
            "workspace_path": str(workspace),
            "template_path": str(template_path),
        },
    )

    assert response.status_code == 401


def test_openapi_exposes_phase2_response_schemas() -> None:
    """All six routes advertise concrete result models in OpenAPI."""
    paths = app.openapi()["paths"]
    expected = {
        "/api/v1/office/word/analyze-template": "WordTemplateAnalysis",
        "/api/v1/office/word/fill-template": "WordTemplateFillResult",
        "/api/v1/office/pdf/read": "PdfReadResult",
        "/api/v1/office/pdf/generate": "PdfGenerateResult",
        "/api/v1/office/pdf/read-form": "PdfFormReadResult",
        "/api/v1/office/pdf/fill-form": "PdfFormFillResult",
    }

    for path, schema_name in expected.items():
        schema = paths[path]["post"]["responses"]["200"]["content"][
            "application/json"
        ]["schema"]
        assert schema["$ref"].endswith(f"/schemas/{schema_name}")
