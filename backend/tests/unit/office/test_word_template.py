"""Unit tests for Word template analysis."""

from datetime import date
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches

from backend.office import word_template
from backend.office.errors import (
    OfficeFileNotFoundError,
    OfficePathError,
    OfficeSizeLimitError,
    OfficeTemplateFillError,
    OfficeTemplateParseError,
)
from backend.office.models import (
    PlaceholderLocation,
    TemplatePlaceholderType,
    WordTemplateFillRequest,
)
from backend.office.word_template import analyze_word_template, fill_word_template


@pytest.fixture()
def simple_template(tmp_path: Path) -> Path:
    """Create a simple Word template with placeholders."""
    doc = Document()
    doc.add_heading("合同模板", level=0)
    doc.add_paragraph("甲方：{{甲方姓名}}")
    doc.add_paragraph("乙方：{{乙方姓名}}")
    doc.add_paragraph("合同金额：{{合同金额}}")

    template_path = tmp_path / "template.docx"
    doc.save(str(template_path))
    return template_path


def test_analyze_simple_template(simple_template: Path):
    result = analyze_word_template(
        simple_template,
        workspace_path=str(simple_template.parent),
    )
    assert len(result.placeholders) == 3
    names = [p.name for p in result.placeholders]
    assert "甲方姓名" in names
    assert "乙方姓名" in names
    assert "合同金额" in names


def test_analyze_template_placeholder_details(simple_template: Path):
    result = analyze_word_template(
        simple_template,
        workspace_path=str(simple_template.parent),
    )
    for ph in result.placeholders:
        assert ph.location == PlaceholderLocation.BODY
        assert ph.type == TemplatePlaceholderType.TEXT
        assert ph.paragraph_index is not None
        assert ph.raw_tag.startswith("{{")
        assert ph.raw_tag.endswith("}}")


def test_analyze_nonexistent_file():
    with pytest.raises(OfficeFileNotFoundError):
        analyze_word_template(
            Path("/nonexistent/template.docx"),
            workspace_path="/tmp",
        )


def test_analyze_invalid_docx(tmp_path: Path):
    invalid_file = tmp_path / "invalid.docx"
    invalid_file.write_text("not a docx file")

    with pytest.raises(OfficeTemplateParseError):
        analyze_word_template(
            invalid_file,
            workspace_path=str(tmp_path),
        )


def test_analyze_non_body_placeholders_and_indices(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("正文")
    body_table = doc.add_table(rows=1, cols=2)
    body_table.cell(0, 0).text = "日期：{{合同日期}}"
    body_table.cell(0, 1).text = "图片：{{签名图片}}"
    nested = body_table.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "嵌套：{{nested_text}}"

    section = doc.sections[0]
    section.header.paragraphs[0].text = "页眉：{{页眉文本}}"
    header_table = section.header.add_table(rows=1, cols=1, width=Inches(1))
    header_table.cell(0, 0).text = "页眉表格：{{header_image}}"
    header_nested = header_table.cell(0, 0).add_table(rows=1, cols=1)
    header_nested.cell(0, 0).text = "页眉嵌套：{{header_nested}}"
    section.footer.paragraphs[0].text = "页脚：{{页脚文本}}"
    footer_table = section.footer.add_table(rows=1, cols=1, width=Inches(1))
    footer_table.cell(0, 0).text = "页脚表格：{{footer_date}}"
    footer_nested = footer_table.cell(0, 0).add_table(rows=1, cols=1)
    footer_nested.cell(0, 0).text = "页脚嵌套：{{footer_nested}}"

    path = tmp_path / "locations.docx"
    doc.save(str(path))

    result = analyze_word_template(path, workspace_path=str(tmp_path))
    placeholders = {placeholder.name: placeholder for placeholder in result.placeholders}

    assert placeholders["合同日期"].type == TemplatePlaceholderType.DATE
    assert placeholders["合同日期"].location == PlaceholderLocation.TABLE
    assert (placeholders["合同日期"].table_index, placeholders["合同日期"].row_index,
            placeholders["合同日期"].col_index) == (0, 0, 0)
    assert placeholders["签名图片"].type == TemplatePlaceholderType.IMAGE
    assert placeholders["nested_text"].location == PlaceholderLocation.TABLE
    assert (placeholders["nested_text"].table_index, placeholders["nested_text"].row_index,
            placeholders["nested_text"].col_index) == (0, 0, 0)
    assert placeholders["页眉文本"].location == PlaceholderLocation.HEADER
    assert placeholders["页脚文本"].location == PlaceholderLocation.FOOTER
    assert placeholders["header_image"].location == PlaceholderLocation.HEADER
    assert (placeholders["header_image"].table_index, placeholders["header_image"].row_index,
            placeholders["header_image"].col_index) == (0, 0, 0)
    assert placeholders["header_nested"].location == PlaceholderLocation.HEADER
    assert (placeholders["header_nested"].table_index, placeholders["header_nested"].row_index,
            placeholders["header_nested"].col_index) == (0, 0, 0)
    assert placeholders["footer_date"].location == PlaceholderLocation.FOOTER
    assert (placeholders["footer_date"].table_index, placeholders["footer_date"].row_index,
            placeholders["footer_date"].col_index) == (0, 0, 0)
    assert placeholders["footer_date"].type == TemplatePlaceholderType.DATE
    assert placeholders["footer_nested"].location == PlaceholderLocation.FOOTER
    assert (placeholders["footer_nested"].table_index, placeholders["footer_nested"].row_index,
            placeholders["footer_nested"].col_index) == (0, 0, 0)


@pytest.mark.parametrize("control_location", ["body_table", "header", "footer"])
def test_analyze_jinja_control_in_all_stories(tmp_path: Path, control_location: str):
    doc = Document()
    if control_location == "body_table":
        outer = doc.add_table(rows=1, cols=1)
        outer.cell(0, 0).add_table(rows=1, cols=1).cell(0, 0).text = "{%tr for row in rows %}"
    elif control_location == "header":
        header_table = doc.sections[0].header.add_table(rows=1, cols=1, width=Inches(1))
        header_table.cell(0, 0).add_table(rows=1, cols=1).cell(0, 0).text = "{% if show_header %}"
    else:
        footer_table = doc.sections[0].footer.add_table(rows=1, cols=1, width=Inches(1))
        footer_table.cell(0, 0).add_table(rows=1, cols=1).cell(0, 0).text = "{% if show_footer %}"

    path = tmp_path / (control_location + ".docx")
    doc.save(str(path))

    result = analyze_word_template(path, workspace_path=str(tmp_path))

    assert result.has_jinja_control is True


def test_analyze_rejects_template_outside_workspace(tmp_path: Path):
    outside = tmp_path.parent / "outside-template.docx"
    Document().save(str(outside))

    with pytest.raises(OfficePathError):
        analyze_word_template(outside, workspace_path=str(tmp_path))


def test_fill_simple_template(simple_template: Path, tmp_path: Path):
    req = WordTemplateFillRequest(
        workspace_path=str(simple_template.parent),
        template_path=str(simple_template),
        output_filename="filled.docx",
        data={"甲方姓名": "张三", "乙方姓名": "李四", "合同金额": "100,000"},
    )
    result = fill_word_template(req)

    assert result.filename == "filled.docx"
    assert result.filled_count == 3
    assert len(result.unfilled_placeholders) == 0
    assert Path(result.output_path).exists()
    filled_doc = Document(result.output_path)
    filled_text = "\n".join(paragraph.text for paragraph in filled_doc.paragraphs)
    assert "甲方：张三" in filled_text
    assert "乙方：李四" in filled_text
    assert "合同金额：100,000" in filled_text


def test_fill_template_partial_data(simple_template: Path, tmp_path: Path):
    req = WordTemplateFillRequest(
        workspace_path=str(simple_template.parent),
        template_path=str(simple_template),
        output_filename="partial.docx",
        data={"甲方姓名": "张三"},
    )
    result = fill_word_template(req)

    assert result.filled_count == 1
    assert len(result.unfilled_placeholders) == 2
    assert "乙方姓名" in result.unfilled_placeholders


def test_fill_rejects_output_path_traversal(simple_template: Path):
    req = WordTemplateFillRequest(
        workspace_path=str(simple_template.parent),
        template_path=str(simple_template),
        output_filename="../filled.docx",
        data={},
    )

    with pytest.raises(OfficeTemplateFillError):
        fill_word_template(req)


def test_fill_rejects_output_path_boundaries(simple_template: Path, tmp_path: Path):
    for output_filename in ("", "/tmp/absolute.docx", "../escape.docx"):
        req = WordTemplateFillRequest(
            workspace_path=str(simple_template.parent),
            template_path=str(simple_template),
            output_filename=output_filename,
            data={},
        )
        with pytest.raises(OfficeTemplateFillError):
            fill_word_template(req)

    outside = tmp_path.parent / "outside-output.docx"
    outside.write_bytes(b"existing")
    symlink = simple_template.parent / "link.docx"
    symlink.symlink_to(outside)
    req = WordTemplateFillRequest(
        workspace_path=str(simple_template.parent),
        template_path=str(simple_template),
        output_filename="link.docx",
        data={},
    )
    with pytest.raises(OfficeTemplateFillError):
        fill_word_template(req)


def test_fill_rejects_dangerous_jinja_expressions(tmp_path: Path):
    for expression in ("{{ ''.__class__.__mro__ }}", "{% import 'os' as os %}", "{% include 'evil.txt' %}"):
        template = tmp_path / "dangerous.docx"
        doc = Document()
        doc.add_paragraph(expression)
        doc.save(str(template))
        req = WordTemplateFillRequest(
            workspace_path=str(tmp_path),
            template_path=str(template),
            output_filename="output.docx",
            data={},
        )
        with pytest.raises(OfficeTemplateFillError):
            fill_word_template(req)
        assert not (tmp_path / "output.docx").exists()


def test_fill_date_and_image_placeholders(tmp_path: Path, fixture_dir: Path):
    image_path = fixture_dir / "signature.png"
    image_path.write_bytes(
        __import__("base64").b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
    )
    doc = Document()
    doc.add_paragraph("日期：{{合同日期}}")
    doc.add_paragraph("签名：{{签名图片}}")
    template = tmp_path / "image-date.docx"
    doc.save(str(template))
    req = WordTemplateFillRequest(
        workspace_path=str(tmp_path),
        template_path=str(template),
        output_filename="image-date-filled.docx",
        data={"合同日期": date(2026, 9, 5)},
        images={"签名图片": str(image_path)},
    )
    result = fill_word_template(req)
    output = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in output.paragraphs)
    assert "2026-09-05" in text
    assert len(output.inline_shapes) == 1


def test_fill_wraps_render_errors(simple_template: Path, monkeypatch):
    class BrokenTemplate:
        def __init__(self, path):
            pass

        def render(self, context):
            raise RuntimeError("render exploded")

    monkeypatch.setattr(word_template, "DocxTemplate", BrokenTemplate)
    req = WordTemplateFillRequest(
        workspace_path=str(simple_template.parent),
        template_path=str(simple_template),
        output_filename="broken.docx",
        data={},
    )

    with pytest.raises(OfficeTemplateFillError, match="Template fill failed"):
        fill_word_template(req)


def test_analyze_rejects_oversized_docx_zip(tmp_path: Path, monkeypatch):
    path = tmp_path / "template.docx"
    Document().save(str(path))
    monkeypatch.setattr(word_template, "MAX_DOCX_COMPRESSED_SIZE", 1)

    with pytest.raises(OfficeSizeLimitError):
        analyze_word_template(path, workspace_path=str(tmp_path))


def test_analyze_rejects_too_many_zip_members(tmp_path: Path, monkeypatch):
    path = tmp_path / "template.docx"
    Document().save(str(path))
    monkeypatch.setattr(word_template, "MAX_DOCX_MEMBERS", 1)

    with pytest.raises(OfficeSizeLimitError):
        analyze_word_template(path, workspace_path=str(tmp_path))


def test_analyze_rejects_too_much_uncompressed_zip_data(tmp_path: Path, monkeypatch):
    path = tmp_path / "template.docx"
    Document().save(str(path))
    monkeypatch.setattr(word_template, "MAX_DOCX_UNCOMPRESSED_SIZE", 1)

    with pytest.raises(OfficeSizeLimitError):
        analyze_word_template(path, workspace_path=str(tmp_path))
