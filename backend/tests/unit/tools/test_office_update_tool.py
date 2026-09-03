# ruff: noqa: UP006, UP007, UP035, UP045 — release/win7 Python 3.8 兼容，保留 typing 注解
"""Unit tests for :mod:`backend.tools.office_update_tool`.

Covers:
- ``requires_tool_context = False`` + WRITE_LOCAL risk declaration
- doc_id 模式：绑定工作区内成功编辑 + DB 状态翻转为 EDITED；未知/越界 id
  折叠为 document_not_found；缺 tool_context → missing_tool_context
- file_path 模式：绝对路径成功编辑；相对路径 / 非 office 扩展名 / 不存在
  文件被拒；ops 非法被拒
- 结果中不泄漏绝对 workspace 路径（doc_id 模式）
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional
from unittest.mock import patch

import pytest

from backend.data.database import Database
from backend.domain.tool_policy import ToolPolicy
from backend.office.models import (
    OfficeDocStatus,
    OfficeDocType,
    OfficeDocumentMetadata,
    OfficeDocumentSummary,
)
from backend.office.session_workspace import bind_session_workspace
from backend.office.storage import save_document
from backend.tools.context import ToolExecutionContext, reset_tool_context, set_tool_context
from backend.tools.office_update_tool import OfficeUpdateTool

pytestmark = pytest.mark.unit


def _tool(**policy_kwargs) -> OfficeUpdateTool:
    return OfficeUpdateTool(policy=ToolPolicy(**policy_kwargs))


def _make_doc(
    *,
    doc_id: str,
    workspace_path: str,
    doc_type: OfficeDocType = OfficeDocType.WORD,
    generated_filename: Optional[str] = None,
) -> OfficeDocumentSummary:
    return OfficeDocumentSummary(
        id=doc_id,
        workspace_path=workspace_path,
        doc_type=doc_type,
        original_filename="上传.docx",
        generated_filename=generated_filename or f"{doc_id}.docx",
        status=OfficeDocStatus.PARSED,
        created_at=1_700_000_000_000,
        updated_at=1_700_000_000_000,
        metadata=OfficeDocumentMetadata(file_size_bytes=1024),
    )


def _seed_session(conn, session_id: str) -> None:
    conn.execute(
        "INSERT INTO sessions (id, title, created_at, updated_at) VALUES (?, ?, ?, ?)",
        (session_id, "t", 1, 1),
    )
    conn.commit()


def _ctx(session_id: str, binding_generation: int = 1) -> ToolExecutionContext:
    return ToolExecutionContext(
        session_id=session_id,
        stream_id="stream-x",
        binding_generation=binding_generation,
        office_doc_scope=frozenset(),
    )


# ── 形状声明 ──────────────────────────────────────────────────────────


def test_tool_declares_write_local_and_no_context_requirement():
    tool = _tool()
    assert tool.requires_tool_context is False
    from backend.domain.risk import RiskClass

    assert tool.risk is RiskClass.WRITE_LOCAL


def test_schema_has_no_workspace_path_parameter():
    props = _tool().schema.parameters["properties"]
    assert "workspace_path" not in props
    assert set(props) == {"doc_id", "file_path", "ops"}


# ── 参数校验 ──────────────────────────────────────────────────────────


def test_missing_ops_rejected():
    assert _tool().execute(doc_id="d1").success is False


def test_ops_must_be_nonempty_dict_list():
    assert _tool().execute(doc_id="d1", ops=[]).success is False
    assert _tool().execute(doc_id="d1", ops=[{"no_op_key": 1}]).success is False


def test_requires_doc_id_or_file_path():
    result = _tool().execute(ops=[{"op": "replace_text", "find": "a", "replace": "b"}])
    assert result.success is False
    assert result.error == "doc_id_or_file_path_required"


# ── doc_id 模式 ───────────────────────────────────────────────────────


def test_update_by_doc_id_edits_file_and_marks_edited(tmp_path: Path):
    db = Database(db_path=str(tmp_path / "t.db"))
    db.init_db()
    conn = db.get_connection()
    _seed_session(conn, "sess-1")
    work = tmp_path / "work"
    work.mkdir()
    binding = bind_session_workspace(conn, "sess-1", str(work), now_ms=1)
    from docx import Document

    doc_file = work / "office" / "word" / "doc-a"
    doc_file.mkdir(parents=True)
    d = Document()
    d.add_paragraph("旧文本")
    d.save(str(doc_file / "doc-a.docx"))
    save_document(conn, _make_doc(doc_id="doc-a", workspace_path=binding.workspace_path))

    with patch("backend.tools.office_update_tool.get_database", return_value=db):
        token = set_tool_context(_ctx("sess-1", binding.generation))
        try:
            result = _tool().execute(
                doc_id="doc-a", ops=[{"op": "replace_text", "find": "旧文本", "replace": "新文本"}]
            )
        finally:
            reset_tool_context(token)

    assert result.success is True
    assert result.content["document_id"] == "doc-a"
    assert result.content["results"][0]["replacements"] == 1
    assert "workspace" not in str(result.content)
    from backend.office.word import read_docx

    parsed = read_docx(doc_file / "doc-a.docx", workspace_path="")
    assert [p.text for p in parsed.paragraphs] == ["新文本"]
    # DB 状态翻转为 EDITED
    row = conn.execute("SELECT status FROM office_documents WHERE id = 'doc-a'").fetchone()
    assert row["status"] == OfficeDocStatus.EDITED.value


def test_update_unknown_doc_is_indistinguishable_not_found(tmp_path: Path):
    db = Database(db_path=str(tmp_path / "t.db"))
    db.init_db()
    conn = db.get_connection()
    _seed_session(conn, "sess-1")
    work = tmp_path / "work"
    work.mkdir()
    binding = bind_session_workspace(conn, "sess-1", str(work), now_ms=1)

    with patch("backend.tools.office_update_tool.get_database", return_value=db):
        token = set_tool_context(_ctx("sess-1", binding.generation))
        try:
            result = _tool().execute(
                doc_id="ghost", ops=[{"op": "replace_text", "find": "a", "replace": "b"}]
            )
        finally:
            reset_tool_context(token)
    assert result.success is False
    assert result.error == "document_not_found"


def test_update_without_context_fails_closed():
    result = _tool().execute(
        doc_id="doc-a", ops=[{"op": "replace_text", "find": "a", "replace": "b"}]
    )
    assert result.success is False
    assert result.error == "missing_tool_context"


def test_update_failed_ops_report_results(tmp_path: Path):
    db = Database(db_path=str(tmp_path / "t.db"))
    db.init_db()
    conn = db.get_connection()
    _seed_session(conn, "sess-1")
    work = tmp_path / "work"
    work.mkdir()
    binding = bind_session_workspace(conn, "sess-1", str(work), now_ms=1)
    from docx import Document

    doc_file = work / "office" / "word" / "doc-b"
    doc_file.mkdir(parents=True)
    d = Document()
    d.add_paragraph("内容")
    d.save(str(doc_file / "doc-b.docx"))
    save_document(conn, _make_doc(doc_id="doc-b", workspace_path=binding.workspace_path))

    with patch("backend.tools.office_update_tool.get_database", return_value=db):
        token = set_tool_context(_ctx("sess-1", binding.generation))
        try:
            result = _tool().execute(
                doc_id="doc-b", ops=[{"op": "replace_text", "find": "不存在", "replace": "x"}]
            )
        finally:
            reset_tool_context(token)
    assert result.success is False
    assert result.error == "operation_failed"
    assert result.content["results"][0]["ok"] is False
    # 文件未被改动
    from backend.office.word import read_docx

    assert "内容" in [
        p.text for p in read_docx(doc_file / "doc-b.docx", workspace_path="").paragraphs
    ]


# ── file_path 模式 ────────────────────────────────────────────────────


def test_update_by_absolute_path(tmp_path: Path):
    from openpyxl import Workbook

    target = tmp_path / "报表.xlsx"
    wb = Workbook()
    wb.active.title = "数据"
    wb.save(str(target))
    result = _tool().execute(
        file_path=str(target),
        ops=[{"op": "append_rows", "sheet": "数据", "rows": [["a", "1"]]}],
    )
    assert result.success is True
    from backend.office.excel import read_xlsx

    assert read_xlsx(target, workspace_path="").sheets[0].rows == [["a", "1"]]


def test_update_by_path_rejects_relative_path():
    result = _tool().execute(
        file_path="报表.xlsx", ops=[{"op": "append_rows", "sheet": "s", "rows": [[1]]}]
    )
    assert result.success is False
    assert result.error.startswith("file_path_absolute_required")


def test_update_by_path_rejects_non_office_extension(tmp_path: Path):
    target = tmp_path / "note.txt"
    target.write_text("hello")
    result = _tool().execute(
        file_path=str(target), ops=[{"op": "replace_text", "find": "h", "replace": "H"}]
    )
    assert result.success is False
    assert result.error.startswith("unsupported_file_type")


def test_update_by_path_missing_file(tmp_path: Path):
    result = _tool().execute(
        file_path=str(tmp_path / "ghost.docx"),
        ops=[{"op": "replace_text", "find": "a", "replace": "b"}],
    )
    assert result.success is False
    assert result.error == "file_not_found"


def test_update_by_path_outside_workspace_root_rejected(tmp_path: Path):
    """hex 链（policy.workspace_root 绑定）下越界 file_path 直接拒绝。"""
    target = tmp_path / "outside.docx"
    from docx import Document

    Document().save(str(target))
    tool = OfficeUpdateTool(policy=ToolPolicy(workspace_root=str(tmp_path / "workspace")))
    result = tool.execute(
        file_path=str(target), ops=[{"op": "replace_text", "find": "a", "replace": "b"}]
    )
    assert result.success is False
    assert result.error.startswith("path_outside_workspace")
