"""Word template analysis and fill using docxtpl.

docxtpl extends python-docx with Jinja2-like template syntax:
- {{variable}} for text placeholders
- {% if condition %} ... {% endif %} for control flow
- {%tr for row in rows %} ... {%tr endfor %} for table loops
"""

from __future__ import annotations

import base64
import re
import time
import zipfile
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docxtpl import DocxTemplate, InlineImage
from jinja2.sandbox import SandboxedEnvironment

from .errors import (
    OfficeFileNotFoundError,
    OfficePathError,
    OfficeSizeLimitError,
    OfficeTemplateFillError,
    OfficeTemplateParseError,
)
from .models import (
    OfficeDocStatus,
    OfficeDocType,
    OfficeDocumentMetadata,
    OfficeDocumentSummary,
    PlaceholderLocation,
    TemplatePlaceholder,
    TemplatePlaceholderType,
    WordTemplateAnalysis,
    WordTemplateFillRequest,
    WordTemplateFillResult,
)
from .path_safety import resolve_within
from .storage import validate_workspace

PLACEHOLDER_RE = re.compile(r"\{\{([^}]+)\}\}")
JINJA_CONTROL_RE = re.compile(r"\{%[^%]+%\}")
MAX_DOCX_COMPRESSED_SIZE = 50 * 1024 * 1024
MAX_DOCX_MEMBERS = 10_000
MAX_DOCX_UNCOMPRESSED_SIZE = 250 * 1024 * 1024
_MAX_IMAGE_BYTES = 10 * 1024 * 1024
_DANGEROUS_TEMPLATE_RE = re.compile(
    r"(?:\{%\s*(?:import|include)\b|__class__|__subclasses__|__mro__|__globals__|__builtins__)",
    re.IGNORECASE,
)


def _validate_docx_zip(file_path: Path) -> None:
    """Reject oversized or malformed DOCX ZIP containers before parsing."""
    try:
        compressed_size = file_path.stat().st_size
        if compressed_size > MAX_DOCX_COMPRESSED_SIZE:
            raise OfficeSizeLimitError(
                compressed_size,
                MAX_DOCX_COMPRESSED_SIZE,
                file_path=file_path,
            )
        with zipfile.ZipFile(str(file_path)) as archive:
            members = archive.infolist()
            if len(members) > MAX_DOCX_MEMBERS:
                raise OfficeSizeLimitError(
                    len(members),
                    MAX_DOCX_MEMBERS,
                    file_path=file_path,
                )
            uncompressed_size = sum(member.file_size for member in members)
            if uncompressed_size > MAX_DOCX_UNCOMPRESSED_SIZE:
                raise OfficeSizeLimitError(
                    uncompressed_size,
                    MAX_DOCX_UNCOMPRESSED_SIZE,
                    file_path=file_path,
                )
    except OfficeSizeLimitError:
        raise
    except (OSError, zipfile.BadZipFile) as exc:
        raise OfficeTemplateParseError(
            "Failed to inspect DOCX ZIP", file_path=file_path
        ) from exc


def _extract_placeholders_from_text(
    text: str,
    location: PlaceholderLocation,
    *,
    paragraph_index: Optional[int] = None,
    table_index: Optional[int] = None,
    row_index: Optional[int] = None,
    col_index: Optional[int] = None,
) -> List[TemplatePlaceholder]:
    """Extract {{}} placeholders from a text string."""
    placeholders = []
    for match in PLACEHOLDER_RE.finditer(text):
        raw_tag = match.group(0)
        name = match.group(1).strip()

        if "日期" in name or "date" in name.lower():
            ph_type = TemplatePlaceholderType.DATE
        elif "图片" in name or "image" in name.lower():
            ph_type = TemplatePlaceholderType.IMAGE
        else:
            ph_type = TemplatePlaceholderType.TEXT

        placeholders.append(
            TemplatePlaceholder(
                name=name,
                raw_tag=raw_tag,
                type=ph_type,
                location=location,
                paragraph_index=paragraph_index,
                table_index=table_index,
                row_index=row_index,
                col_index=col_index,
            )
        )
    return placeholders


def _scan_paragraphs(doc: Document) -> List[TemplatePlaceholder]:
    """Scan all body paragraphs for placeholders."""
    placeholders = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        if PLACEHOLDER_RE.search(text):
            placeholders.extend(
                _extract_placeholders_from_text(
                    text, PlaceholderLocation.BODY, paragraph_index=idx
                )
            )
    return placeholders


def _cell_direct_text(cell) -> str:
    """Return paragraph text directly in a cell, excluding nested tables."""
    return "\n".join(paragraph.text for paragraph in cell.paragraphs)


def _scan_table(table, location: PlaceholderLocation, table_index: int) -> List[TemplatePlaceholder]:
    """Scan a table and nested tables using the outer table index."""
    placeholders = []
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            text = _cell_direct_text(cell)
            if PLACEHOLDER_RE.search(text):
                placeholders.extend(
                    _extract_placeholders_from_text(
                        text,
                        location,
                        table_index=table_index,
                        row_index=row_idx,
                        col_index=col_idx,
                    )
                )
            for nested_table in cell.tables:
                placeholders.extend(_scan_table(nested_table, location, table_index))
    return placeholders


def _scan_tables(doc: Document) -> List[TemplatePlaceholder]:
    """Scan all tables, including nested tables."""
    placeholders = []
    for table_idx, table in enumerate(doc.tables):
        placeholders.extend(_scan_table(table, PlaceholderLocation.TABLE, table_idx))
    return placeholders


def _scan_story_tables(story, location: PlaceholderLocation) -> List[TemplatePlaceholder]:
    """Scan tables in a document story, including nested tables."""
    placeholders = []
    for table_idx, table in enumerate(story.tables):
        placeholders.extend(_scan_table(table, location, table_idx))
    return placeholders


def _scan_headers_footers(doc: Document) -> List[TemplatePlaceholder]:
    """Scan paragraphs and tables in headers and footers."""
    placeholders = []
    for section in doc.sections:
        for story, location in (
            (section.header, PlaceholderLocation.HEADER),
            (section.footer, PlaceholderLocation.FOOTER),
        ):
            for para in story.paragraphs:
                if PLACEHOLDER_RE.search(para.text):
                    placeholders.extend(
                        _extract_placeholders_from_text(para.text, location)
                    )
            placeholders.extend(_scan_story_tables(story, location))
    return placeholders


def _table_has_jinja_control(table) -> bool:
    """Check direct cell text and nested tables for Jinja controls."""
    for row in table.rows:
        for cell in row.cells:
            if JINJA_CONTROL_RE.search(_cell_direct_text(cell)):
                return True
            if any(_table_has_jinja_control(nested) for nested in cell.tables):
                return True
    return False


def _story_has_jinja_control(story) -> bool:
    """Check paragraphs and nested table cells in one document story."""
    if any(JINJA_CONTROL_RE.search(para.text) for para in story.paragraphs):
        return True
    return any(_table_has_jinja_control(table) for table in story.tables)


def _has_jinja_control(doc: Document) -> bool:
    """Check body, tables, headers, and footers for Jinja2 controls."""
    if _story_has_jinja_control(doc):
        return True
    return any(
        _story_has_jinja_control(story)
        for section in doc.sections
        for story in (section.header, section.footer)
    )


def _build_template_summary(
    file_path: Path,
    *,
    document_id: str,
    workspace_path: str,
) -> OfficeDocumentSummary:
    """Build a summary for the template analysis result."""
    now_ms = int(time.time() * 1000)
    return OfficeDocumentSummary(
        id=document_id,
        workspace_path=workspace_path,
        doc_type=OfficeDocType.WORD,
        original_filename=file_path.name,
        generated_filename=file_path.name,
        status=OfficeDocStatus.PARSED,
        created_at=now_ms,
        updated_at=now_ms,
        metadata=OfficeDocumentMetadata(file_size_bytes=file_path.stat().st_size),
    )


def analyze_word_template(
    file_path: Path,
    *,
    workspace_path: str,
    document_id: Optional[str] = None,
) -> WordTemplateAnalysis:
    """Analyze a Word template and extract all {{}} placeholders."""
    file_path = Path(file_path)
    workspace = validate_workspace(Path(workspace_path))
    if not file_path.exists():
        raise OfficeFileNotFoundError(file_path)
    file_path = resolve_within(workspace, file_path)

    if not file_path.exists():
        raise OfficeFileNotFoundError(file_path)
    if not file_path.is_file():
        raise OfficeTemplateParseError(
            "Path is not a file", file_path=file_path
        )

    _validate_docx_zip(file_path)
    try:
        doc = Document(str(file_path))
    except Exception as exc:
        raise OfficeTemplateParseError(
            "Failed to parse DOCX", file_path=file_path
        ) from exc

    placeholders: List[TemplatePlaceholder] = []
    placeholders.extend(_scan_paragraphs(doc))
    placeholders.extend(_scan_tables(doc))
    placeholders.extend(_scan_headers_footers(doc))

    doc_id = document_id or file_path.stem
    summary = _build_template_summary(
        file_path, document_id=doc_id, workspace_path=workspace_path
    )

    return WordTemplateAnalysis(
        file_path=str(file_path),
        placeholders=placeholders,
        summary=summary,
        has_jinja_control=_has_jinja_control(doc),
    )


def _validate_template_safety(file_path: Path) -> None:
    """Reject template constructs that could access files or Python internals."""
    try:
        with zipfile.ZipFile(str(file_path)) as archive:
            xml_members = (
                member for member in archive.infolist() if member.filename.endswith(".xml")
            )
            for member in xml_members:
                if _DANGEROUS_TEMPLATE_RE.search(archive.read(member).decode("utf-8", "ignore")):
                    raise OfficeTemplateFillError("Template contains unsafe expressions")
    except OfficeTemplateFillError:
        raise
    except Exception as exc:
        raise OfficeTemplateFillError("Template safety check failed") from exc


def _normalize_template_value(value: Any) -> Any:
    if isinstance(value, datetime):
        return value.isoformat(sep=" ")
    if isinstance(value, date):
        return value.isoformat()
    return value


def _load_image_descriptor(value: str, workspace: Path) -> BytesIO:
    if value.startswith("data:image/"):
        try:
            _, encoded = value.split(",", 1)
            payload = base64.b64decode(encoded, validate=True)
        except (ValueError, base64.binascii.Error) as exc:
            raise OfficeTemplateFillError("Invalid image data") from exc
    else:
        image_path = resolve_within(workspace, Path(value))
        if not image_path.is_file():
            raise OfficeTemplateFillError("Image file is unavailable")
        try:
            payload = image_path.read_bytes()
        except OSError as exc:
            raise OfficeTemplateFillError("Image file cannot be read") from exc
    if len(payload) > _MAX_IMAGE_BYTES:
        raise OfficeTemplateFillError("Image exceeds size limit")
    return BytesIO(payload)


def _build_fill_context(
    template: DocxTemplate,
    req: WordTemplateFillRequest,
    placeholder_names: set,
    workspace: Path,
) -> Dict[str, Any]:
    context = {
        name: _normalize_template_value(req.data.get(name, ""))
        for name in placeholder_names
    }
    for name, descriptor in (req.images or {}).items():
        context[name] = InlineImage(template, _load_image_descriptor(descriptor, workspace))
    return context


def fill_word_template(req: WordTemplateFillRequest) -> WordTemplateFillResult:
    """Fill a Word template with a restricted docxtpl environment."""
    template_path = Path(req.template_path)
    analysis = analyze_word_template(template_path, workspace_path=req.workspace_path)
    template_path = Path(analysis.file_path)
    workspace = validate_workspace(Path(req.workspace_path))
    _validate_template_safety(template_path)

    output_name = req.output_filename
    if not output_name or Path(output_name).is_absolute():
        raise OfficeTemplateFillError("Invalid output filename")
    if "/" in output_name or "\\" in output_name:
        raise OfficeTemplateFillError("Invalid output filename")
    try:
        output_path = resolve_within(workspace, template_path.parent / output_name)
    except OfficePathError as exc:
        raise OfficeTemplateFillError("Invalid output filename") from exc
    if output_path == template_path or output_path.is_dir():
        raise OfficeTemplateFillError("Invalid output filename")
    if output_path.exists():
        raise OfficeTemplateFillError("Invalid output filename")

    placeholder_names = {placeholder.name for placeholder in analysis.placeholders}
    provided_names = set(req.data.keys()) | set((req.images or {}).keys())
    unfilled = sorted(placeholder_names - provided_names)
    try:
        template = DocxTemplate(str(template_path))
        context = _build_fill_context(template, req, placeholder_names, workspace)
        template.render(context, jinja_env=SandboxedEnvironment(autoescape=False))
        template.save(str(output_path))
        return WordTemplateFillResult(
            output_path=str(output_path),
            filename=output_name,
            file_size_bytes=output_path.stat().st_size,
            filled_count=len(placeholder_names) - len(unfilled),
            unfilled_placeholders=unfilled,
        )
    except OfficeTemplateFillError:
        raise
    except Exception as exc:
        raise OfficeTemplateFillError("Template fill failed") from exc
