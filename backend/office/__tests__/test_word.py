"""Tests for backend.office.word (DOCX reader + generator).

TDD approach (plan §4.1.2): these tests are written FIRST. Each fixture is
programmatically generated via python-docx in the test body itself.

Test coverage (plan §5.1):
- empty DOCX (no paragraphs)
- DOCX with heading levels (h1, h2, h3)
- DOCX with body paragraphs (Normal style)
- DOCX with a table
- DOCX with an image
- Missing file → OfficeFileNotFoundError
- Corrupt file → OfficeParseError
"""
from __future__ import annotations

import base64
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches

from backend.office.errors import OfficeFileNotFoundError, OfficeParseError
from backend.office.word import read_docx


def _make_minimal_png(path: Path) -> Path:
    """1x1 red PNG for image-embedding tests."""
    b64 = (
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4nGP8z8Dw"
        "HwAFBQIAX8v0gQAAAABJRU5ErkJggg=="
    )
    path.write_bytes(base64.b64decode(b64, validate=True))
    return path


def _build_empty_docx(path: Path) -> Path:
    """DOCX with no body content."""
    Document().save(path)
    return path


def _build_headings_docx(path: Path) -> Path:
    """DOCX with h1, h2, h3 + body paragraphs."""
    doc = Document()
    doc.add_heading("Heading One", level=1)
    doc.add_paragraph("Body text under H1.")
    doc.add_heading("Heading Two", level=2)
    doc.add_paragraph("Body text under H2.")
    doc.add_heading("Heading Three", level=3)
    doc.add_paragraph("Body text under H3.")
    doc.add_paragraph("Plain trailing paragraph.")
    doc.save(path)
    return path


def _build_table_docx(path: Path) -> Path:
    """DOCX with a 2x3 table."""
    doc = Document()
    doc.add_paragraph("Before table")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(0, 2).text = "C"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    table.cell(1, 2).text = "3"
    doc.add_paragraph("After table")
    doc.save(path)
    return path


def _build_image_docx(path: Path, fixture_dir: Path) -> Path:
    """DOCX with an embedded image."""
    img = _make_minimal_png(fixture_dir / "tiny.png")
    doc = Document()
    doc.add_paragraph("Document with image:")
    doc.add_picture(str(img), width=Inches(1))
    doc.save(path)
    return path


# ──────────────────────────────────────────────────────────────────────
# read_docx tests
# ──────────────────────────────────────────────────────────────────────


def test_read_docx_returns_read_result_shape(fixture_dir: Path) -> None:
    """Smoke test: returns OfficeWordReadResult with summary + paragraphs + tables."""
    path = _build_headings_docx(fixture_dir / "headings.docx")
    result = read_docx(path)
    assert result.summary.doc_type.value == "word"
    assert result.summary.metadata.file_size_bytes > 0
    assert isinstance(result.paragraphs, list)
    assert isinstance(result.tables, list)


def test_read_docx_empty_document_returns_zero_paragraphs(fixture_dir: Path) -> None:
    """Empty DOCX: paragraphs is empty list."""
    path = _build_empty_docx(fixture_dir / "empty.docx")
    result = read_docx(path)
    assert result.paragraphs == []
    assert result.tables == []
    assert result.images == 0


def test_read_docx_extracts_heading_levels(fixture_dir: Path) -> None:
    """Heading paragraphs: level field matches add_heading(level=N)."""
    path = _build_headings_docx(fixture_dir / "headings.docx")
    result = read_docx(path)

    # Find the heading paragraphs
    h1 = next(p for p in result.paragraphs if "Heading One" in p.text)
    h2 = next(p for p in result.paragraphs if "Heading Two" in p.text)
    h3 = next(p for p in result.paragraphs if "Heading Three" in p.text)
    assert h1.level == 1
    assert h2.level == 2
    assert h3.level == 3

    # Body paragraphs have level 0
    body = next(p for p in result.paragraphs if "Plain trailing" in p.text)
    assert body.level == 0


def test_read_docx_extracts_paragraph_text(fixture_dir: Path) -> None:
    """Body paragraphs: text field matches add_paragraph(...).text."""
    path = _build_headings_docx(fixture_dir / "headings.docx")
    result = read_docx(path)

    texts = [p.text for p in result.paragraphs]
    assert "Heading One" in texts
    assert "Body text under H1." in texts
    assert "Plain trailing paragraph." in texts


def test_read_docx_extracts_table_data(fixture_dir: Path) -> None:
    """Tables in document body are NOT in result.paragraphs; they're in result.tables."""
    path = _build_table_docx(fixture_dir / "table.docx")
    result = read_docx(path)

    # One table extracted
    assert len(result.tables) == 1
    table = result.tables[0]
    assert table.rows == [["A", "B", "C"], ["1", "2", "3"]]

    # Surrounding paragraphs are in result.paragraphs
    texts = [p.text for p in result.paragraphs]
    assert "Before table" in texts
    assert "After table" in texts


def test_read_docx_counts_images(fixture_dir: Path) -> None:
    """DOCX with one image: result.images == 1."""
    path = _build_image_docx(fixture_dir / "image.docx", fixture_dir)
    result = read_docx(path)
    assert result.images == 1


def test_read_docx_missing_file_raises_file_not_found(fixture_dir: Path) -> None:
    """Nonexistent path raises OfficeFileNotFoundError."""
    missing = fixture_dir / "does-not-exist.docx"
    with pytest.raises(OfficeFileNotFoundError):
        read_docx(missing)


def test_read_docx_corrupt_file_raises_parse_error(fixture_dir: Path) -> None:
    """Non-DOCX bytes raise OfficeParseError."""
    garbage = fixture_dir / "garbage.docx"
    garbage.write_bytes(b"definitely not a docx file")
    with pytest.raises(OfficeParseError):
        read_docx(garbage)


def test_read_docx_metadata_paragraph_count(fixture_dir: Path) -> None:
    """Summary.metadata.paragraph_count matches result.paragraphs length."""
    path = _build_headings_docx(fixture_dir / "headings.docx")
    result = read_docx(path)
    assert result.summary.metadata.paragraph_count == len(result.paragraphs)
    # Fixture adds 7 non-empty paragraphs: 3 headings + 3 body under each + 1 trailing
    assert result.summary.metadata.paragraph_count == 7
