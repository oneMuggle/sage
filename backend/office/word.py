"""DOCX reader using python-docx.

Pure functions: no FastAPI, no I/O outside the file argument. Caller (the
FastAPI route handler in office_routes.py) wraps exceptions into HTTP errors.

The reader extracts a structured view of a .docx:
- paragraphs with style + level (heading level 0 means body text)
- tables (rows of cell text, ignoring nested tables / images inside cells)
- image count (inline shapes count as images)

It does NOT extract:
- headers / footers (out of scope per plan §1.3)
- footnotes / endnotes
- text boxes / shapes outside the body
- track changes / comments

These omissions are intentional per plan §1.3 "non-goals".
"""

from __future__ import annotations

import time
from pathlib import Path
from typing import List, Optional

from docx import Document

from .errors import OfficeFileNotFoundError, OfficeParseError
from .models import (
    OfficeDocStatus,
    OfficeDocType,
    OfficeDocumentMetadata,
    OfficeDocumentSummary,
    OfficeWordReadResult,
    WordParagraphContent,
    WordTableContent,
)


def _extract_heading_level(style_name: str) -> int:
    """Extract heading level from python-docx style name.

    Returns 1, 2, 3, ... for 'Heading 1', 'Heading 2', etc.
    Returns 0 for 'Normal' or any other body style.
    """
    if not style_name:
        return 0
    if style_name.startswith("Heading "):
        try:
            return int(style_name[len("Heading ") :])
        except (ValueError, IndexError):
            return 0
    if style_name == "Title":
        return 1  # treat Title as h1
    return 0


def _extract_paragraphs(doc: Document) -> List[WordParagraphContent]:
    """Extract body paragraphs (skipping tables, headers, footers)."""
    paragraphs: List[WordParagraphContent] = []
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()
        # Skip empty paragraphs (would just clutter result)
        if not text:
            continue
        paragraphs.append(
            WordParagraphContent(
                style=style_name,
                text=text,
                level=_extract_heading_level(style_name),
            )
        )
    return paragraphs


def _extract_tables(doc: Document) -> List[WordTableContent]:
    """Extract body tables (skipping nested tables)."""
    tables: List[WordTableContent] = []
    for table in doc.tables:
        rows: List[List[str]] = []
        for row in table.rows:
            cells: List[str] = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(cells)
        tables.append(WordTableContent(rows=rows))
    return tables


def _count_images(doc: Document) -> int:
    """Count inline pictures in the document body.

    python-docx exposes doc.inline_shapes as a sequence of InlineShape objects,
    one per picture (image or chart).
    """
    return len(doc.inline_shapes)


def _build_docx_summary(
    file_path: Path,
    *,
    document_id: str,
    workspace_path: str,
    generated_filename: Optional[str],
    original_filename: Optional[str],
    status: OfficeDocStatus,
    paragraph_count: int,
) -> OfficeDocumentSummary:
    """Construct a summary for a DOCX document."""
    now_ms = int(time.time() * 1000)
    return OfficeDocumentSummary(
        id=document_id,
        workspace_path=workspace_path,
        doc_type=OfficeDocType.WORD,
        original_filename=original_filename,
        generated_filename=generated_filename or file_path.name,
        status=status,
        created_at=now_ms,
        updated_at=now_ms,
        metadata=OfficeDocumentMetadata(
            page_count=None,  # DOCX doesn't expose page count without rendering
            table_count=None,
            paragraph_count=paragraph_count,
            file_size_bytes=file_path.stat().st_size,
        ),
    )


def read_docx(
    file_path: Path,
    *,
    document_id: Optional[str] = None,
    workspace_path: str = "",
    generated_filename: Optional[str] = None,
    original_filename: Optional[str] = None,
) -> OfficeWordReadResult:
    """Read a .docx file and return its structured content.

    Args:
        file_path: Absolute path to the .docx file.
        document_id: Optional UUID for the summary record. Defaults to file_path.stem.
        workspace_path: Required by storage layer; pass empty string for read-only tests.
        generated_filename: Filename as stored in workspace/office/<id>/.
        original_filename: User's uploaded filename.

    Returns:
        OfficeWordReadResult with summary + paragraphs + tables + image count.

    Raises:
        OfficeFileNotFoundError: file doesn't exist.
        OfficeParseError: file exists but isn't a valid DOCX.
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise OfficeFileNotFoundError(file_path)
    if not file_path.is_file():
        raise OfficeParseError(f"Path is not a regular file: {file_path}", file_path=file_path)

    try:
        doc = Document(str(file_path))
    except Exception as exc:
        # python-docx raises zipfile.BadZipFile, lxml.etree.XMLSyntaxError, etc.
        raise OfficeParseError(f"Failed to parse DOCX: {exc}", file_path=file_path) from exc

    paragraphs = _extract_paragraphs(doc)
    tables = _extract_tables(doc)
    images = _count_images(doc)

    doc_id = document_id or file_path.stem

    summary = _build_docx_summary(
        file_path,
        document_id=doc_id,
        workspace_path=workspace_path,
        generated_filename=generated_filename,
        original_filename=original_filename,
        status=OfficeDocStatus.PARSED,
        paragraph_count=len(paragraphs),
    )

    return OfficeWordReadResult(
        summary=summary,
        paragraphs=paragraphs,
        tables=tables,
        images=images,
    )


# ──────────────────────────────────────────────────────────────────────
# Generator (Phase 1.4 step 19, plan §4.1.4)
# ──────────────────────────────────────────────────────────────────────


def generate_docx(req) -> Path:
    """Generate a .docx file from structured Pydantic input."""
    import uuid

    from docx import Document as _Doc

    from .errors import OfficeGenerateError
    from .models import OfficeDocType
    from .path_safety import managed_document_path
    from .storage import validate_workspace

    workspace = validate_workspace(Path(req.workspace_path))
    doc_id = uuid.uuid4().hex
    # Compose the full file path with one validated call. Raises
    # OfficePathError on filename separators, parent traversal, wrong
    # extension, doc_id injection, or any path that escapes the workspace.
    output_path = managed_document_path(workspace, OfficeDocType.WORD, doc_id, req.filename)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        doc = _Doc()
        # Title
        doc.add_heading(req.title, level=0)
        # Body paragraphs
        for para in req.paragraphs:
            if para.heading == "h1":
                doc.add_heading(para.text, level=1)
            elif para.heading == "h2":
                doc.add_heading(para.text, level=2)
            elif para.heading == "h3":
                doc.add_heading(para.text, level=3)
            else:
                doc.add_paragraph(para.text)
        # Tables
        for table_spec in req.tables:
            table = doc.add_table(rows=1 + len(table_spec.rows), cols=len(table_spec.headers))
            # Header row
            for ci, header in enumerate(table_spec.headers):
                table.cell(0, ci).text = header
            # Data rows
            for ri, row in enumerate(table_spec.rows):
                for ci, cell in enumerate(row):
                    if ci < len(table_spec.headers):
                        table.cell(ri + 1, ci).text = cell
        doc.save(str(output_path))
    except Exception as exc:
        raise OfficeGenerateError(f"Failed to generate DOCX: {exc}", file_path=output_path) from exc

    return output_path
